"""
生成 v1.4.2 设置界面 UI 设计审查报告 (DOCX)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE, "UI设计审查报告_v1.4.2.docx")

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

# 标题样式
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Microsoft YaHei'
    h.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1E)

def set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ClaudeFloat 浮窗工具\n设置界面 UI 设计审查报告")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("版本 v1.4.2  |  2026-07-14\n基于 claude_floating_launcher.py (L346–L850)")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 概述
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. 审查概述", level=1)
doc.add_paragraph(
    "本报告对 ClaudeFloat v1.4.2 设置对话框 (SettingsDialog) 的 UI 设计进行系统性审查。"
    "审查范围涵盖字体层级、颜色系统、间距与布局、交互反馈、可访问性五个维度。"
    "每个问题附改进建议及多个备选方案，供后续版本迭代参考。"
)

doc.add_paragraph(
    "审查对象：SettingsDialog (QDialog, FramelessWindowHint + WA_TranslucentBackground)\n"
    "窗口策略：420px 最小宽度，高度由内容自动计算（~700–920px）\n"
    "布局结构：外层 8px margin → container (圆角 16px) → VBoxLayout (间距 12px)"
)

# ══════════════════════════════════════════════════════════════
# 2. 当前设计元素清单
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. 当前设计元素清单", level=1)

doc.add_heading("2.1 字体层级", level=2)
add_table(doc,
    ["元素", "字号", "字重", "用途", "代码位置"],
    [
        ["对话框标题", "15pt", "Bold", '"Claude Code 浮窗设置"', "_label(size=15, bold=True) L470"],
        ["版本号", "9pt", "Normal", '"v1.4.2"', "_label(size=9) L473"],
        ["关闭按钮 ✕", "14pt", "Normal", "标题栏关闭按钮", "self._font(14) L482"],
        ["GroupBox 标题", "12pt", "Bold", "各设置模块标题", "box.setFont(self._font(12, bold=True))"],
        ["Radio 选项文字", "11pt", "Normal", "普通模式 / 跳过权限 / 主题", "rb.setFont(self._font(11)) L511"],
        ["Checkbox 选项文字", "11pt", "Normal", "边缘吸附 / 自动隐藏 / 清理进程", "cb.setFont(self._font(11))"],
        ["Radio CSS 字号", "未指定", "—", "QRadioButton 样式表中无 font-size", "仅 color+spacing L394"],
        ["Checkbox CSS 字号", "11px", "—", "QCheckBox 样式表指定", "font-size: 11px L400"],
        ["普通按钮", "12px", "Normal", "浏览/重置/取消/应用", "btn CSS font-size: 12px L385"],
        ["保存按钮", "12px", "Bold", "主要操作按钮", "save_btn CSS font-size: 12px L390"],
        ["Slider 数值标签", "12pt", "Bold", '"52 px" / "88%"', "_label(size=12, bold=True)"],
        ["Slider 描述标签", "12pt", "Normal", '"边长:" / "不透明度:"', "_label(size=12)"],
        ["LineEdit 文字", "10pt", "Normal", "工作目录路径", "self._font(10) L669"],
        ["安全警告标签", "9pt", "Bold", "跳过权限安全提示", "self._font(9, bold=True) L552"],
        ["Log 路径提示", "8pt", "Normal", "底部日志路径", "_label(size=8) L717"],
        ["对话框默认字体", "9pt", "Normal", "setFont(self._font(9)) L447", "（极少被继承）"],
    ]
)

doc.add_heading("2.2 颜色系统", level=2)
doc.add_paragraph(
    "采用 iOS 风格双主题配色 (THEMES dict L126–147)，通过 get_colors() 返回字典。"
    "主要语义色：Accent = iOS 蓝 #007AFF / #0A84FF。"
    "二级文字色 text_secondary 硬编码在 _build_styles() 中（非来自 THEMES dict）："
    "亮色 #3C3C43，暗色 #EBEBF5。"
)

add_table(doc,
    ["颜色角色", "亮色值", "暗色值", "CSS 变量 / 使用位置"],
    [
        ["主背景 (SURFACE)", "#F2F2F7", "#2C2C2E", "container bg, hover bg"],
        ["主文字 (TEXT)", "#1C1C1E", "#F2F2F7", "GroupBox 标题, LineEdit 文字"],
        ["次要文字 (HINT)", "#8E8E93", "#98989D", "版本号, 日志路径, 关闭按钮"],
        ["卡片白 (card_bg)", "#FFFFFF", "#2C2C2E", "Radio/Checkbox indicator bg, 按钮 bg"],
        ["二级文字 (text_secondary)", "#3C3C43", "#EBEBF5", "Radio/Checkbox 文字 (硬编码)"],
        ["强调色 (ACCENT)", "#007AFF", "#0A84FF", "indicator checked, slider, 保存按钮"],
        ["分隔线 (SEPARATOR)", "#E5E5EA", "#38383A", "标题栏下方分隔线, slider groove"],
        ["边框 (BORDER)", "#FFFFFF", "#48484A", "container 边框, indicator 未选中边框"],
        ["安全警告文字", "#FF3B30", "#FF3B30", "warn_label: 亮/暗均红色"],
        ["安全警告背景", "#FFE5E5", "#FFE5E5", "warn_label: 亮/暗均浅红"],
    ]
)

doc.add_heading("2.3 间距与布局", level=2)
add_table(doc,
    ["参数", "值", "说明"],
    [
        ["Container 圆角", "16px", "settingsContainer border-radius"],
        ["Container 内边距", "20,12,20,16 (L,R,T,B)", "layout.setContentsMargins"],
        ["外层 margin", "8px (四边)", "outer.setContentsMargins (阴影留白)"],
        ["模块间间距", "12px", "layout.setSpacing(12)"],
        ["GroupBox 内间距", "6px / 8px", "vl.setSpacing(6 or 8)"],
        ["按钮圆角", "8px", "btn CSS border-radius"],
        ["Indicator 圆角", "Radio 10px / Checkbox 6px", "iOS 圆环风格"],
        ["Indicator 尺寸", "20×20px", "Radio/Checkbox 统一"],
        ["Radio spacing", "8px", "CSS spacing: 8px (indicator↔text)"],
        ["主按钮内边距", "8px 20px", "保存按钮 padding"],
        ["普通按钮内边距", "8px 16px", "浏览/取消/应用 padding"],
        ["LineEdit 内边距", "6px 8px", "目录路径输入框"],
        ["关闭按钮尺寸", "28×28px", "标题栏 ✕ 按钮"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 3. 问题诊断
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. 问题诊断", level=1)

# 3.1
doc.add_heading("3.1 字体层级混乱（严重程度：高）", level=2)
doc.add_paragraph(
    "当前设置对话框存在至少 8 种不同字号（8pt, 9pt, 10pt, 11pt, 11px(≈8.25pt), 12pt, 14pt, 15pt），"
    "但缺乏清晰的层级逻辑："
)
doc.add_paragraph(
    "• Radio 文字同时受 QFont(11pt) 和 CSS font-size 双重影响。Checkbox CSS 指定 font-size:11px "
    "而 Radio CSS 未指定，导致不同类型选项文字渲染大小不一致（11px ≈ 8.25pt vs 11pt）。\n"
    "• 默认对话框字体设为 9pt (line 447) 但几乎所有子控件显式覆盖，此值成为死设置。\n"
    "• Slider 描述标签 (\"边长:\", \"不透明度:\") 使用 12pt，与 GroupBox 标题 (12pt Bold) 同号但语义层级完全不同。",
    style='List Bullet'
)

# 3.2
doc.add_heading("3.2 Radio / Checkbox 字号不一致（严重程度：中）", level=2)
doc.add_paragraph(
    "Radio 按钮：通过 rb.setFont(self._font(11)) 设置 QFont 11pt，但 CSS 中未指定 font-size。\n"
    "Checkbox 按钮：CSS 中设置了 font-size: 11px，同时 cb.setFont(self._font(11)) 设置 QFont 11pt。\n"
    "Qt 的 CSS font-size 使用 px 单位，而 QFont 使用 pt 单位 (1pt ≈ 1.333px)。结果：\n"
    "• Radio 实际渲染 ~14.7px（11pt × 1.333）\n"
    "• Checkbox 实际渲染 11px（CSS 覆盖了 QFont）\n"
    "这导致同一对话框中 Radio 文字明显大于 Checkbox 文字，视觉不统一。"
)

# 3.3
doc.add_heading("3.3 二级文字色硬编码（严重程度：中）", level=2)
doc.add_paragraph(
    "text_secondary 颜色值在 _build_styles() (line 376) 中硬编码为 \"#3C3C43\"(亮) / \"#EBEBF5\"(暗)，"
    "而非来自 THEMES 字典。这意味着：\n"
    "• 如果未来新增主题（如 Sepia、High Contrast），需额外修改 _build_styles()\n"
    "• 与 THEMES 中的 HINT 色用途边界模糊：HINT=#8E8E93 vs text_secondary=#3C3C43，"
    "但\"次要文字\"这个概念被两个不同色值服务"
)

# 3.4
doc.add_heading("3.4 GroupBox 标题样式过于朴素（严重程度：低）", level=2)
doc.add_paragraph(
    "GroupBox CSS 仅设置 color + border:none + padding-top:12px，无字体规格、无左侧色条、"
    "无与下方内容的视觉分组（背景色无差异）。各模块视觉区分仅靠 12px 间距。"
    "在较长的设置列表中（当前 7 个模块），用户扫读时难以快速定位目标模块。"
)

# 3.5
doc.add_heading("3.5 按钮层次扁平（严重程度：低）", level=2)
doc.add_paragraph(
    "底部 4 个操作按钮（应用、取消、保存）中仅\"保存\"使用 accent 实色填充，其余均为白底蓝字。\n"
    "但\"应用\"和\"保存\"同为确认类操作，视觉权重差异巨大但语义相近，可能造成混淆。\n"
    "\"应用\"按钮的功能（关闭对话框 + 预览设置）与标签\"应用\"不够匹配。"
)

# 3.6
doc.add_heading("3.6 安全警告暗色模式可读性（严重程度：中）", level=2)
doc.add_paragraph(
    "warn_label 使用固定的红色文字 (#FF3B30) + 浅红色背景 (#FFE5E5)，在暗色模式下：\n"
    "• 浅红色背景 (#FFE5E5) 在暗色卡片 (#2C2C2E) 上形成高对比度矩形，视觉突兀\n"
    "• 红色文字在浅红背景上可读性尚可，但暗色模式下通常期望深色背景 + 低饱和警告色"
)

# 3.7
doc.add_heading("3.7 缺少视觉焦点指示（严重程度：中）", level=2)
doc.add_paragraph(
    "Radio/Checkbox/Button 在 Tab 键导航时无 focus 样式（CSS 中未定义 :focus 伪状态）。\n"
    "键盘用户无法判断当前焦点位置。这是可访问性 (a11y) 的基础缺失。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. 改进建议与备选方案
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. 改进建议与备选方案", level=1)

# ── 4.1 字体层级重整 ──
doc.add_heading("4.1 字体层级重整 (对应 3.1, 3.2)", level=2)
doc.add_paragraph("目标：建立清晰的 4 级字体层级，消除 Radio/Checkbox 字号不一致。")

doc.add_heading("方案 A — 严格 4 级层级（推荐）", level=3)
add_table(doc,
    ["层级", "字号", "字重", "适用于"],
    [
        ["H1 — 页面标题", "16pt (≈21px)", "Bold", "对话框标题"],
        ["H2 — 模块标题", "13pt (≈17px)", "Bold", "GroupBox 标题"],
        ["H3 — 选项文字", "11pt (≈15px)", "Normal", "Radio, Checkbox 标签"],
        ["Body — 辅助信息", "9pt (≈12px)", "Normal", "版本号, 日志路径, 描述文字"],
    ]
)
doc.add_paragraph(
    "• 统一 Radio 和 Checkbox 的字号设置方式：移除 CSS 中 font-size，统一使用 QFont(11pt)\n"
    "• 移除默认对话框字体 setFont(9)（已是死设置）\n"
    "• Slider 标签 ('边长:', '不透明度:') 降为 Body 级 (9pt)，数值标签 ('52 px') 保持 H3 级"
)

doc.add_heading("方案 B — iOS HIG 风格", level=3)
add_table(doc,
    ["层级", "字号", "字重", "适用于"],
    [
        ["Large Title", "17pt", "Bold", "对话框标题"],
        ["Section Header", "13pt", "Semibold", "GroupBox 标题 + 左侧色条"],
        ["Body", "11pt", "Normal", "Radio, Checkbox, 标签"],
        ["Caption", "8pt", "Normal", "版本号, 提示文字"],
    ]
)
doc.add_paragraph(
    "• 仅 4 种字号，严格对齐 iOS 设置应用的视觉节奏\n"
    "• Section Header 使用全大写 + 字间距增加，强调模块边界"
)

doc.add_heading("方案 C — 保守微调（最小改动）", level=3)
doc.add_paragraph(
    "• 仅修复 Radio/Checkbox 字号不一致（统一为 11pt QFont，移除 CSS font-size）\n"
    "• GroupBox 标题保持 12pt Bold，其余不变\n"
    "• 优点：风险最低，仅需改 2 行代码"
)

# ── 4.2 颜色系统规范化 ──
doc.add_heading("4.2 颜色系统规范化 (对应 3.3)", level=2)
doc.add_paragraph("目标：将 text_secondary 纳入 THEMES 字典，支持未来多主题扩展。")

doc.add_heading("方案 A — 扩展 THEMES 字典（推荐）", level=3)
doc.add_paragraph(
    "在 THEMES['light'] / THEMES['dark'] 中增加 'TEXT_SECONDARY' 键：\n"
    "  light: (60, 60, 67)   — #3C3C43\n"
    "  dark:  (235, 235, 245) — #EBEBF5\n"
    "_build_styles() 中从 self._c 读取而非硬编码。"
)
doc.add_paragraph(
    "• 优点：一处定义，全局使用；未来新增主题只需添加 THEMES 条目\n"
    "• 缺点：需修改 THEMES 结构和 get_colors() 调用处"
)

doc.add_heading("方案 B — 复用 HINT 色（零改动）", level=3)
doc.add_paragraph(
    "• 直接使用 HINT 色（#8E8E93）替代 text_secondary\n"
    "• 优点：无需编码\n"
    "• 缺点：HINT (142,142,147) 比当前 text_secondary (60,60,67) 对比度更低，可能影响可读性"
)

# ── 4.3 模块视觉强化 ──
doc.add_heading("4.3 模块视觉强化 (对应 3.4)", level=2)
doc.add_paragraph("目标：增强 GroupBox 标题的视觉存在感，帮助用户快速定位设置模块。")

doc.add_heading("方案 A — 左侧 accent 色条（推荐）", level=3)
doc.add_paragraph(
    "在 GroupBox 标题左侧添加 3px 宽的 accent 色条，通过 CSS border-left 或 QSS 实现：\n"
    "  QGroupBox::title { border-left: 3px solid #007AFF; padding-left: 8px; }"
)
doc.add_paragraph(
    "• 优点：符合现代 UI 趋势（Settings app, VS Code），视觉锚点明确\n"
    "• 缺点：Qt CSS 对 ::title 子控件的支持有限，可能需要自定义 QGroupBox 子类"
)

doc.add_heading("方案 B — 卡片式分隔", level=3)
doc.add_paragraph(
    "为每个 GroupBox 添加微妙的背景色（如 SURFACE 色 + 5% 亮度）和独立的圆角边框：\n"
    "  QGroupBox { background: rgba(242,242,247,0.5); border-radius: 10px; padding: 12px; }"
)
doc.add_paragraph(
    "• 优点：模块间视觉隔离更明显，扫读效率高\n"
    "• 缺点：增加垂直空间占用，每个卡片的内边距累加"
)

doc.add_heading("方案 C — 分割线分隔", level=3)
doc.add_paragraph(
    "在各模块之间添加 1px 分隔线（与标题栏下方 sep 风格一致），替代依赖间距分隔：\n"
    "• 优点：实现简单，与现有 sep 组件复用\n"
    "• 缺点：分隔线 + 标题 = 双重视觉元素，可能显重"
)

# ── 4.4 按钮层次优化 ──
doc.add_heading("4.4 按钮层次优化 (对应 3.5)", level=2)
doc.add_paragraph("目标：使底部操作按钮的功能更直观，视觉层次更合理。")

doc.add_heading("方案 A — 三档按钮（推荐）", level=3)
add_table(doc,
    ["按钮", "样式", "语义"],
    [
        ["保存", "Accent 实色填充 (Primary)", "确认并持久化所有更改"],
        ["应用", "Accent 边框 + 透明背景 (Secondary)", "即时应用但不持久化"],
        ["取消", "灰色文字 + 透明背景 (Tertiary)", "放弃更改并关闭"],
    ]
)
doc.add_paragraph(
    "• \"应用\"改为不关闭对话框，仅调用 parent.apply_settings()\n"
    "• \"保存\"同时关闭对话框\n"
    "• 按钮顺序改为 macOS 风格：[取消] [应用] [保存]（确认性操作在右）"
)

doc.add_heading("方案 B — 简化双按钮", level=3)
doc.add_paragraph(
    "• 移除\"应用\"按钮，仅保留 [取消] [保存]\n"
    "• 主题切换已即时生效（v1.4.2），其余设置通过\"保存\"一次性确认\n"
    "• 优点：减少用户决策负担，按钮更简洁"
)

# ── 4.5 安全警告暗色适配 ──
doc.add_heading("4.5 安全警告暗色模式适配 (对应 3.6)", level=2)
doc.add_paragraph("目标：warn_label 在暗色模式下不显得突兀。")

doc.add_heading("方案 A — 暗色专用警告色（推荐）", level=3)
doc.add_paragraph(
    "在 THEMES 中增加 warn_bg / warn_fg 键：\n"
    "  light: warn_bg=#FFE5E5, warn_fg=#FF3B30\n"
    "  dark:  warn_bg=#3D1F1F, warn_fg=#FF6B6B (更低饱和的背景 + 稍亮的文字)"
)
doc.add_paragraph(
    "• 优点：暗色下警告边框/背景与暗色卡片融合，不刺眼\n"
    "• 缺点：需增加 THEMES 键，CSS 在主题切换时也需动态更新 warn_label"
)

doc.add_heading("方案 B — 半透明叠加", level=3)
doc.add_paragraph(
    "warn_label 使用半透明红色背景 (rgba(255,59,48,0.1)) + 白色文字：\n"
    "• 优点：自动适配任意主题背景色\n"
    "• 缺点：Qt CSS 对 rgba() 支持有限（Qt 5.x 不完全支持 CSS3 rgba）"
)

# ── 4.6 焦点指示器 ──
doc.add_heading("4.6 键盘焦点指示器 (对应 3.7)", level=2)
doc.add_paragraph("目标：Tab 键导航时提供清晰的视觉焦点反馈。")

doc.add_heading("方案 A — CSS :focus 轮廓（推荐）", level=3)
doc.add_paragraph(
    "为 Radio/Checkbox/Button/Slider 添加 :focus 伪状态样式：\n"
    "  QRadioButton:focus { outline: 2px solid #007AFF; outline-offset: 2px; }\n"
    "  QCheckBox:focus { outline: 2px solid #007AFF; outline-offset: 2px; }\n"
    "  QPushButton:focus { border: 2px solid #007AFF; }"
)
doc.add_paragraph(
    "• 优点：标准做法，所有平台通用\n"
    "• 缺点：Qt 的 outline 支持有限，可能需要改用 border 模拟"
)

doc.add_heading("方案 B — accent 色描边", level=3)
doc.add_paragraph(
    "使用与 accent 同色的虚线边框表示焦点：\n"
    "  :focus { border: 1px dashed #007AFF; }"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. 优先级矩阵
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. 改进优先级矩阵", level=1)

add_table(doc,
    ["优先级", "问题", "影响", "改动量", "建议方案"],
    [
        ["P0 — 立即", "3.2 Radio/Checkbox 字号不一致", "视觉不统一，同一对话框文字大小不同", "2 行", "4.1 方案 C — 统一用 QFont"],
        ["P1 — 本版本", "3.1 字体层级混乱", "缺乏设计系统性", "~15 行", "4.1 方案 A — 4 级层级"],
        ["P1 — 本版本", "3.7 缺少焦点指示器", "键盘不可用 (a11y)", "~10 行", "4.6 方案 A — :focus outline"],
        ["P2 — 下版本", "3.6 安全警告暗色适配", "暗色模式视觉突兀", "~5 行", "4.5 方案 A — 暗色警告色"],
        ["P2 — 下版本", "3.3 text_secondary 硬编码", "扩展性差", "~8 行", "4.2 方案 A — 纳入 THEMES"],
        ["P3 — 后续", "3.4 GroupBox 标题朴素", "扫读效率", "~20 行", "4.3 方案 A — 左侧色条"],
        ["P3 — 后续", "3.5 按钮层次扁平", "操作语义混淆", "~10 行", "4.4 方案 A — 三档按钮"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 6. 附录
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. 附录：推荐设计令牌 (Design Tokens)", level=1)

doc.add_paragraph(
    "如采纳 4.1 方案 A + 4.2 方案 A，建议在代码中定义如下令牌结构："
)

doc.add_paragraph(
    "# 字体令牌\n"
    "FONT_H1 = (16, Bold)      # 页面标题\n"
    "FONT_H2 = (13, Bold)      # 模块标题\n"
    "FONT_H3 = (11, Normal)    # 选项文字 (Radio/Checkbox)\n"
    "FONT_BODY = (9, Normal)   # 辅助信息\n\n"
    "# 间距令牌\n"
    "SPACE_XS = 4px\n"
    "SPACE_SM = 8px\n"
    "SPACE_MD = 12px\n"
    "SPACE_LG = 20px\n\n"
    "# 圆角令牌\n"
    "RADIUS_SM = 6px    (checkbox)\n"
    "RADIUS_MD = 8px    (button)\n"
    "RADIUS_LG = 10px   (radio)\n"
    "RADIUS_XL = 16px   (container)",
    style='List Bullet'
)

# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"报告已生成: {OUTPUT}")
print(f"文件大小: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
